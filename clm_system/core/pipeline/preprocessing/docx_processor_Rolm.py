import docx
import re
import logging
import asyncio
from typing import Dict, Any, List, Tuple
from PIL import Image
import io
import os
import numpy as np
import cv2
import requests
import json
from pathlib import Path

# Set up logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DOCXProcessor:
    """Processes DOCX files into contract JSON structure using RolmOCR for image tables."""
    
    def __init__(self, rolm_api_key=None, rolm_endpoint="http://192.168.10.1:1234/v1"):
        """
        Initialize the DOCX processor.
        
        Args:
            rolm_api_key: API key for RolmOCR service
            rolm_endpoint: Endpoint URL for RolmOCR API
        """
        self.rolm_api_key = "rolmocr"
        self.rolm_endpoint = "http://192.168.10.1:1234/v1/chat/completions"
        
        if not self.rolm_api_key:
            logger.warning("No RolmOCR API key provided. Set ROLM_API_KEY environment variable or pass it to the constructor.")
    
    async def process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Process a DOCX file into a contract structure.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Contract object in dictionary form
        """
        try:
            # Extract text and images from DOCX
            raw_data = await self._extract_content(file_path)
            text = raw_data["text"]
            images = raw_data["images"]
            # Fix: Use raw string to avoid backslash issues
            logger.info(f"Extracted {len(text.split(os.linesep))} lines of text and {len(images)} images")
            
            # Process the text into contract structure
            contract = await self._structure_contract(text, images, file_path)
            
            return contract
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}", exc_info=True)
            raise
    
    async def _extract_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and images from DOCX.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with text and images
        """
        text_lines = []
        images = []  # List of (paragraph_index, image_data) tuples
        try:
            logger.info(f"Opening document: {file_path}")
            doc = docx.Document(file_path)
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    text_lines.append(text)
                
                # Check for images in the paragraph
                for run in para.runs:
                    if hasattr(run, 'element') and run.element.xpath('.//wp:inline'):
                        try:
                            blip = run.element.xpath('.//a:blip')[0]
                            rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            part = doc.part.related_parts[rid]
                            image = Image.open(io.BytesIO(part.blob))
                            images.append((i, image))
                            logger.debug(f"Found image at paragraph {i}")
                        except (IndexError, KeyError) as e:
                            logger.warning(f"Failed to extract image at paragraph {i}: {e}")
            
            # Also check tables for text content
            for i, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(' | '.join(row_text))
                
                if table_text:
                    text_lines.append("".join(table_text))
                    logger.debug(f"Added text from table {i}")
            
            full_text = os.linesep.join(text_lines)
            return {"text": full_text, "images": images}
        except Exception as e:
            logger.error(f"Error extracting content from DOCX: {e}", exc_info=True)
            raise
    
    async def _structure_contract(self, text: str, images: List[Tuple], file_path: str) -> Dict[str, Any]:
        """
        Convert raw text and images into contract structure.
        
        Args:
            text: Extracted text from the document
            images: List of (paragraph_index, image) tuples
            file_path: Original file path
            
        Returns:
            Structured contract as a dictionary
        """
        lines = text.split(os.linesep)
        header = {}
        sections = []
        current_section = None
        # Improved regex to match section titles by number patterns
        section_pattern = re.compile(r'^\d+\.(?:\d+\.)*\s*(?:[A-Z\s/]+|\w+\s*[\w\s]*(?:M3|M³).*|\w.*)$')
        image_idx = 0
        
        # Process lines
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Check if line is a section title
            if section_pattern.match(line):
                logger.debug(f"Found section: {line}")
                if current_section:
                    sections.append(current_section)
                current_section = {"title": line, "content": []}
            else:
                if current_section:
                    # Add to current section content
                    current_section["content"].append(line)
                else:
                    # Header parsing before first section
                    if ":" in line:
                        key, value = line.split(":", 1)
                        header[key.strip().lower()] = value.strip()
                    else:
                        header.setdefault("notes", []).append(line)
            
            # Handle images based on their position in document
            while image_idx < len(images) and images[image_idx][0] <= i:
                _, image = images[image_idx]
                logger.info(f"Processing table image at position {i}")
                table_data = await self._extract_table_from_image(image)
                
                if current_section:
                    table_entry = {"type": "table", "data": table_data}
                    current_section["content"].append(table_entry)
                else:
                    header.setdefault("tables", []).append(table_data)
                
                image_idx += 1
        
        # Add the last section if it exists
        if current_section:
            sections.append(current_section)
        
        # Structure the contract with dynamic party extraction
        contract = {
            "title": self._determine_title(header, file_path),
            "metadata": {
                "contract_type": self._determine_contract_type(sections, header),
                "parties": self._extract_parties(sections, header),
                "status": "draft",
                "tags": self._extract_tags(sections, header)
            },
            "header": header,
            "clauses": self._process_clauses(sections)
        }
        
        return contract
    
    def _determine_title(self, header: Dict, file_path: str) -> str:
        """Extract title from header or filename."""
        if "title" in header:
            return header["title"]
        elif "re" in header:
            return header["re"]
        else:
            # Fallback to filename
            return Path(file_path).stem.replace("_", " ").title()
    
    def _determine_contract_type(self, sections: List[Dict], header: Dict) -> str:
        """Determine contract type based on content."""
        full_text = " ".join([section.get("title", "") for section in sections])
        full_text = full_text.lower()
        
        if "purchase" in full_text or "sale" in full_text:
            return "purchase"
        elif "agreement" in full_text:
            return "agreement"
        elif "offer" in full_text:
            return "offer"
        else:
            return "contract"
    
    def _extract_tags(self, sections: List[Dict], header: Dict) -> List[str]:
        """Extract relevant tags based on content."""
        # Get all text for analysis
        all_text = " ".join([section.get("title", "") for section in sections])
        all_text += " " + " ".join([str(v) for v in header.values() if isinstance(v, str)])
        all_text = all_text.lower()
        
        tags = []
        if "gasoline" in all_text:
            tags.append("gasoline")
        if "naphtha" in all_text:
            tags.append("naphtha")
        if "oil" in all_text:
            tags.append("oil")
        if "trading" in all_text or "trader" in all_text:
            tags.append("trading")
        
        # Default tags if none found
        if not tags:
            tags = ["contract", "business"]
            
        return tags
    
    def _extract_parties(self, sections: List[Dict], header: Dict) -> List[Dict]:
        """Extract party information from sections or header."""
        parties = []
        buyer = None
        seller = None
        
        # First try to find in sections
        for section in sections:
            title = section["title"].lower()
            content_text = ""
            
            # Collect text content from section
            for item in section["content"]:
                if isinstance(item, str):
                    content_text += item + " "
            
            if ("buyer" in title or "purchaser" in title) and not buyer:
                buyer = {
                    "role": "buyer",
                    "name": self._extract_company_name(content_text) or "Unknown Buyer",
                    "id": "party-001"
                }
            elif ("seller" in title or "vendor" in title) and not seller:
                seller = {
                    "role": "seller",
                    "name": self._extract_company_name(content_text) or "Unknown Seller", 
                    "id": "party-002"
                }
        
        # Fallback to header if sections don't provide party info
        if not buyer and "to" in header:
            buyer = {
                "role": "buyer",
                "name": self._extract_company_name(header["to"]) or "Unknown Buyer",
                "id": "party-001"
            }
            
        if not seller and "from" in header:
            seller = {
                "role": "seller",
                "name": self._extract_company_name(header["from"]) or "TRADECO",
                "id": "party-002"
            }
        
        # Default fallbacks
        if not buyer:
            buyer = {"role": "buyer", "name": "Unknown Buyer", "id": "party-001"}
        if not seller:
            seller = {"role": "seller", "name": "TRADECO", "id": "party-002"}
        
        parties.append(buyer)
        parties.append(seller)
        
        return parties
    
    def _extract_company_name(self, text: str) -> str:
        """Extract company name from text."""
        # Common company suffixes
        company_suffixes = [
            "Inc", "LLC", "Ltd", "Limited", "Corp", "Corporation", 
            "Co", "Company", "GmbH", "AG", "SA", "PLC", "BV", "NV"
        ]
        
        # Try to find company name patterns
        lines = text.split(os.linesep)
        for line in lines:
            line = line.strip()
            if any(suffix in line for suffix in company_suffixes):
                return line
            
            # Look for capitalized words that might be a company name
            words = line.split()
            if len(words) >= 2 and all(word[0].isupper() for word in words if word and word[0].isalpha()):
                return line
        
        # If no clear company pattern, return the first non-empty line
        for line in lines:
            if line.strip():
                return line.strip()
                
        return ""
    
    def _process_clauses(self, sections: List[Dict]) -> List[Dict]:
        """
        Process sections into structured clauses.
        
        Args:
            sections: List of section dictionaries
            
        Returns:
            List of processed clauses
        """
        clauses = []
        
        for section in sections:
            clause = {
                "title": section["title"],
                "content": []
            }
            
            # Process content items with proper structure
            for item in section["content"]:
                if isinstance(item, str):
                    clause["content"].append({"type": "text", "value": item})
                elif isinstance(item, dict):
                    if "type" in item and item["type"] == "table":
                        clause["content"].append(item)
                    else:
                        # Handle legacy format
                        if "table" in item:
                            clause["content"].append({"type": "table", "data": item["table"]})
            
            clauses.append(clause)
        
        return clauses
    
    async def _extract_table_from_image(self, image: Image.Image) -> List[List[str]]:
        """
        Extract text from table images using RolmOCR.
        
        Args:
            image: PIL Image containing a table
            
        Returns:
            Structured table as a list of rows, each row being a list of cell texts
        """
        try:
            # Check if RolmOCR API key is available
            if not self.rolm_api_key:
                logger.warning("No RolmOCR API key available. Falling back to simplified table extraction.")
                return self._fallback_table_extraction(image)
            
            # Prepare image for RolmOCR
            image_bytes = io.BytesIO()
            # Convert to RGB in case it's RGBA to avoid issues
            if image.mode == 'RGBA':
                image = image.convert('RGB')
            image.save(image_bytes, format='JPEG')
            image_bytes.seek(0)
            
            # Prepare request for RolmOCR API
            headers = {
                #"Authorization": f"Bearer {self.rolm_api_key}",
                "Content-Type": "application/json"
            }
            
            # Convert image to base64
            import base64
            image_base64 = base64.b64encode(image_bytes.read()).decode('utf-8')
            
            # Prepare request data
            data = {
                "model": "rolmocr", # Use the model identifier from LM Studio
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            # Prompt asking the model to extract the table
                            {"type": "text", "text": "Extract the table structure and all text content from the following image as accurately as possible. Present the result as rows and columns, perhaps using markdown or a clear delimiter."},
                            # Image data
                            {"type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}
                            }
                        ]
                    }
                ],
                "max_tokens": 2048, # Adjust as needed
                "temperature": 0.1 # Low temperature for factual extraction
            }
            # Send request to RolmOCR
            response = requests.post(
                self.rolm_endpoint,
                headers=headers,
                json=data
            )
            
            if response.status_code != 200:
                logger.error(f"RolmOCR API error: {response.status_code}, {response.text}")
                return self._fallback_table_extraction(image)
            
            # Process response
            ocr_result = response.json()
            
            # Extract tables from response
            if "tables" in ocr_result and ocr_result["tables"]:
                # Process tables from RolmOCR result
                table_data = []
                for table in ocr_result["tables"]:
                    rows = []
                    for row in table["cells"]:
                        row_data = []
                        for cell in row:
                            row_data.append(cell["text"])
                        if row_data:  # Skip empty rows
                            rows.append(row_data)
                    table_data.append(rows)
                
                # Return first table or empty if none found
                return table_data[0] if table_data else []
            elif "text" in ocr_result:
                # Fall back to parsing text as table
                return self._parse_table_data(ocr_result["text"])
            else:
                logger.warning("No tables found in RolmOCR response")
                return self._fallback_table_extraction(image)
                
        except Exception as e:
            logger.error(f"Error extracting table with RolmOCR: {e}", exc_info=True)
            return self._fallback_table_extraction(image)
    
    def _fallback_table_extraction(self, image: Image.Image) -> List[List[str]]:
        """
        Fallback method when RolmOCR is not available or fails.
        Uses OpenCV for basic table structure detection.
        
        Args:
            image: PIL Image containing a table
            
        Returns:
            Placeholder table structure
        """
        try:
            # Convert PIL Image to OpenCV format
            img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Detect lines to find table structure
            edges = cv2.Canny(gray, 50, 150, apertureSize=3)
            lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=100, minLineLength=100, maxLineGap=10)
            
            # Count horizontal and vertical lines to estimate table size
            h_lines = []
            v_lines = []
            
            if lines is not None:
                for line in lines:
                    x1, y1, x2, y2 = line[0]
                    if abs(x2 - x1) > abs(y2 - y1):  # Horizontal line
                        h_lines.append((y1 + y2) // 2)  # Use average y-coordinate
                    else:  # Vertical line
                        v_lines.append((x1 + x2) // 2)  # Use average x-coordinate
            
            # Sort and remove duplicates
            h_lines = sorted(list(set([y for y in h_lines])))
            v_lines = sorted(list(set([x for x in v_lines])))
            
            # Create empty table structure
            rows = max(2, len(h_lines) - 1)  # At least 2 rows
            cols = max(2, len(v_lines) - 1)  # At least 2 columns
            
            # Create placeholder table
            return [["[Cell data not available]" for _ in range(cols)] for _ in range(rows)]
            
        except Exception as e:
            logger.error(f"Fallback table extraction failed: {e}", exc_info=True)
            return [["[Table extraction failed]"]]
    
    def _parse_table_data(self, text: str) -> List[List[str]]:
        """
        Parse text into structured table data.
        
        Args:
            text: Raw OCR text
            
        Returns:
            List of rows, each containing a list of cell values
        """
        # Split rows and clean data
        rows = [row.strip() for row in text.split(os.linesep) if row.strip()]
        
        # Process rows into table structure
        table = []
        for row in rows:
            # Split on multiple spaces to identify columns
            cells = re.split(r'\s{2,}', row)
            # Clean empty cells
            cells = [cell.strip() for cell in cells if cell.strip()]
            
            if cells:  # Skip empty rows
                table.append(cells)
        
        # Table structure validation
        if len(table) < 1:
            return [["Table structure not recognized"]]
        
        # Attempt to align columns based on header/first row
        num_columns = max([len(row) for row in table[:3]], default=0)
        if num_columns > 0:
            aligned_table = []
            for row in table:
                # Pad rows with missing columns
                while len(row) < num_columns:
                    row.append('')
                # Trim excess columns
                aligned_table.append(row[:num_columns])
            return aligned_table
        
        return table


# Example usage
async def process_document(file_path, output_path=None, api_key=None):
    """
    Process a document and save the result as JSON.
    
    Args:
        file_path: Path to the DOCX file
        output_path: Path to save the JSON output (default: based on input filename)
        api_key: RolmOCR API key (optional)
    """
    processor = DOCXProcessor(rolm_api_key=api_key)
    
    # Process document
    logger.info(f"Processing document: {file_path}")
    result = await processor.process_docx(file_path)
    
    # Determine output path
    if not output_path:
        output_path = Path(file_path).with_suffix('.json')
    
    # Save result
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    
    logger.info(f"Processing complete. Output saved to: {output_path}")
    return result

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Process DOCX files into JSON structure")
    parser.add_argument("file_path", help="Path to the DOCX file to process")
    parser.add_argument("--output", "-o", help="Path to save the JSON output")
    parser.add_argument("--api-key", help="RolmOCR API key (can also be set via ROLM_API_KEY env variable)")
    
    args = parser.parse_args()
    
    asyncio.run(process_document(args.file_path, args.output, args.api_key))